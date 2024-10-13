import os
import platform
import subprocess
import sys

# Define common paths and directory structures
base_folder = r'C:\Users\micky\OneDrive\Desktop\XDR' if platform.system() == 'Windows' else os.path.expanduser('~/Desktop/XDR')

# Ensure the folder structure exists
def create_directories():
    if not os.path.exists(base_folder):
        os.makedirs(base_folder)
    print(f"Base XDR directory created at: {base_folder}")

# Check if the current platform is Windows or Unix-based (Linux/macOS)
def check_platform():
    os_platform = platform.system()
    print(f"Detected platform: {os_platform}")
    return os_platform

# Install necessary Python libraries
def install_python_libs():
    required_libs = ['python-docx', 'openpyxl', 'pandas']
    subprocess.check_call([sys.executable, "-m", "pip", "install", *required_libs])
    print("Installed necessary Python libraries.")

# Install external tools like nmap, Wireshark, Volatility for forensic analysis
def install_external_tools(os_platform):
    try:
        if os_platform == 'Windows':
            # Windows specific installations (e.g., Chocolatey for tool installation)
            subprocess.run(['powershell', '-Command', 'Set-ExecutionPolicy RemoteSigned -Scope CurrentUser -Force'])
            subprocess.run(['powershell', '-Command', 'Set-ExecutionPolicy Bypass -Scope Process -Force'])
            subprocess.run(['powershell', '-Command', 'iex', '(new-object net.webclient).DownloadString("https://chocolatey.org/install.ps1")'])
            subprocess.run(['choco', 'install', 'nmap', 'wireshark', 'volatility'])

        elif os_platform in ['Linux', 'Darwin']:  # For Linux or macOS
            subprocess.run(['sudo', 'apt-get', 'update'])
            subprocess.run(['sudo', 'apt-get', 'install', '-y', 'nmap', 'wireshark', 'volatility'])
    except Exception as e:
        print(f"Error during tool installation: {e}")
    print("Installed external forensic and network tools.")

# Placeholder function for forensic data collection (network scans, memory analysis)
def collect_forensic_data():
    # For example, scanning the network with nmap
    print("Running network scan using nmap...")
    try:
        scan_output = subprocess.check_output(['nmap', '-sP', '192.168.1.0/24'])
        print(scan_output.decode())
    except Exception as e:
        print(f"Error running nmap scan: {e}")

    # More forensic data collection can be added here (e.g., memory dumps, file integrity checks)

# Generate a DOCX forensic report
def generate_docx_report(report_name, data):
    from docx import Document

    doc = Document()
    doc.add_heading('XDR Forensic Report', 0)
    
    # Adding content to the DOCX
    doc.add_paragraph(f'Report Name: {report_name}')
    doc.add_paragraph(f'Description: {data["description"]}')
    
    # Adding tables to the report
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Key'
    hdr_cells[1].text = 'Value'
    
    for key, value in data['table_data'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    # Saving the report to the specified folder
    report_path = os.path.join(base_folder, f"{report_name}.docx")
    doc.save(report_path)
    print(f"DOCX report saved to {report_path}")

# Generate an XLSX forensic report
def generate_xlsx_report(report_name, data):
    import pandas as pd

    df = pd.DataFrame(data['table_data'].items(), columns=['Key', 'Value'])
    
    # Saving the dataframe as an Excel file
    report_path = os.path.join(base_folder, f"{report_name}.xlsx")
    df.to_excel(report_path, index=False)
    print(f"XLSX report saved to {report_path}")

# Main function that sets up everything
def setup_xdr():
    os_platform = check_platform()
    
    create_directories()
    
    # Install Python libraries and external tools
    install_python_libs()
    install_external_tools(os_platform)
    
    # Collect forensic data
    collect_forensic_data()

    # Placeholder forensic data for reporting
    forensic_data = {
        'description': 'Unauthorized access detected to sensitive files.',
        'table_data': {
            'IP Address': '192.168.1.1',
            'Affected User': 'JaneDoe',
            'Files Accessed': 'confidential.docx, secret_plans.xlsx',
            'Timestamp': '2024-09-27 14:32:21',
        }
    }

    # Generate reports
    generate_docx_report('Forensic_Report', forensic_data)
    generate_xlsx_report('Forensic_Report', forensic_data)

if __name__ == "__main__":
    setup_xdr()
